import os
import re
import PyPDF2
from docx import Document
from src.clean_resume import clean_resume_text  # your text cleaning function


def pdf_extraction(file_path: str) -> str:
    """Extract text from a PDF file safely."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"PDF file not found: {file_path}")

    text = ""
    try:
        with open(file_path, "rb") as f:
            reader = PyPDF2.PdfReader(f)
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    except Exception as e:
        raise RuntimeError(f"Error reading PDF file: {e}")

    return clean_resume_text(text)


def document_extraction(file_path: str) -> str:
    """Extract text from DOCX file safely, including tables and bullet points."""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"DOCX file not found: {file_path}")

    try:
        doc = Document(file_path)
    except Exception as e:
        raise RuntimeError(f"Error opening DOCX file: {e}")

    raw_lines = []

    # 1. Extract paragraphs
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            text = re.sub(r"\s+", " ", text)
            raw_lines.append(text)

    # 2. Extract tables
    for table in doc.tables:
        for row in table.rows:
            cells = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    cell_text = re.sub(r"\s+", " ", cell_text)
                    cells.append(cell_text)
            if cells:
                raw_lines.append(" | ".join(cells))

    # 3. Merge broken bullet lines
    merged_lines = []
    buffer = ""
    for line in raw_lines:
        if line.startswith(("•", "-", "–")):
            if buffer:
                merged_lines.append(buffer.strip())
            buffer = line
        else:
            if buffer:
                buffer += " " + line
            else:
                merged_lines.append(line)
    if buffer:
        merged_lines.append(buffer.strip())

    # 4. Final cleanup
    final_lines = [re.sub(r"\s+", " ", line) for line in merged_lines]

    return "\n".join(final_lines)


def extract_resume(file_path: str) -> str:
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"Resume file not found: {file_path}")

    ext = os.path.splitext(file_path)[1].lower()

    if ext == ".pdf":
        return pdf_extraction(file_path)
    elif ext in [".docx", ".doc"]:
        return document_extraction(file_path)
    else:
        raise ValueError(f"Unsupported file format: {ext}")